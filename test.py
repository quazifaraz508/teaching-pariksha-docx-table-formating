import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os

def set_cell_borders(cell):
    """
    Set dark borders for a table cell
    """
    tc_pr = cell._element.tcPr
    
    # Define border properties for dark grid lines
    border_attrs = {
        'sz': '4',           # Border size (4 = 1/2 pt)
        'val': 'single',     # Border type
        'color': '000000',   # Black color
        'space': '0'         # No space
    }
    
    # Set borders for all sides
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = tc_pr.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tc_pr.append(border)
        
        for attr, value in border_attrs.items():
            border.set(qn(f'w:{attr}'), value)

def set_table_borders(table):
    """
    Set dark borders for entire table
    """
    tbl_pr = table._element.tblPr
    
    # Create or get table borders element
    tbl_borders = tbl_pr.find(qn('w:tblBorders'))
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    
    # Define border properties
    border_attrs = {
        'sz': '6',           # Slightly thicker borders (6 = 3/4 pt)
        'val': 'single',     # Single line
        'color': '000000',   # Black
        'space': '0'         # No space
    }
    
    # Set borders for all sides including inside borders
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tbl_borders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tbl_borders.append(border)
        
        for attr, value in border_attrs.items():
            border.set(qn(f'w:{attr}'), value)

def fix_table_background(doc):
    """
    Convert tables with black background to normal tables with dark grid lines
    """
    tables_fixed = 0
    
    for table in doc.tables:
        tables_fixed += 1
        # Set table style to have dark grid lines
        table.style = 'Table Grid'
        
        # Remove any shading from the entire table and set dark borders
        for row in table.rows:
            for cell in row.cells:
                # Remove cell shading
                try:
                    tc_pr = cell._element.tcPr
                    shd = tc_pr.find(qn('w:shd'))
                    if shd is not None:
                        shd.set(qn('w:fill'), 'auto')
                        shd.set(qn('w:val'), 'clear')
                except:
                    pass
                
                # Set dark borders for the cell
                set_cell_borders(cell)
                
                # Also check and fix paragraph formatting in cells
                for paragraph in cell.paragraphs:
                    # Reset paragraph formatting
                    try:
                        paragraph.style = doc.styles['Normal']
                    except:
                        pass
                    
                    # Reset run formatting
                    for run in paragraph.runs:
                        try:
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Set text to black
                            run.font.bold = False
                            run.font.italic = False
                        except:
                            pass
    
    return doc, tables_fixed

def fix_table_background_enhanced(doc):
    """
    Enhanced version with better control over table grid appearance
    """
    tables_fixed = 0
    
    for table in doc.tables:
        tables_fixed += 1
        # Remove table style first
        table.style = None
        
        # Apply dark grid borders to entire table
        set_table_borders(table)
        
        for row in table.rows:
            for cell in row.cells:
                # Remove cell shading
                try:
                    tc_pr = cell._element.tcPr
                    shd = tc_pr.find(qn('w:shd'))
                    if shd is not None:
                        shd.set(qn('w:fill'), 'auto')
                        shd.set(qn('w:val'), 'clear')
                except:
                    pass
                
                # Reset text formatting
                for paragraph in cell.paragraphs:
                    try:
                        paragraph.style = doc.styles['Normal']
                    except:
                        pass
                    
                    for run in paragraph.runs:
                        try:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = False
                            run.font.italic = False
                        except:
                            pass
    
    return doc, tables_fixed

def main():
    st.set_page_config(
        page_title="Teaching Pariksha Table Formatter",
        page_icon="📋",
        layout="wide"
    )
    
    st.title("📋Teaching Pariksha Table Formatter")
    st.markdown("""
    This tool helps you fix Word documents with tables that have black backgrounds by converting them 
    to normal tables with dark grid lines and proper text formatting.
    """)
    
    # Sidebar for options
    st.sidebar.header("Settings")
    processing_mode = st.sidebar.radio(
        "Processing Mode:",
        ["Standard"],
        help="Standard: Uses Table Grid style. Enhanced: Custom borders with more control."
    )
    
    # File upload
    uploaded_file = st.file_uploader(
        "Upload your Word document (.docx)", 
        type=['docx'],
        help="Select a Word document with tables to process"
    )
    
    if uploaded_file is not None:
        # Display file info
        file_details = {
            "Filename": uploaded_file.name,
            "File size": f"{uploaded_file.size / 1024:.2f} KB"
        }
        st.write("File details:")
        st.json(file_details)
        
        # Process document
        if st.button("🔄 Process Document", type="primary"):
            try:
                with st.spinner("Processing your document..."):
                    # Read the uploaded file
                    doc = Document(io.BytesIO(uploaded_file.getvalue()))
                    
                    # Process based on selected mode
                    if processing_mode == "Standard":
                        processed_doc, tables_fixed = fix_table_background(doc)
                    else:
                        processed_doc, tables_fixed = fix_table_background_enhanced(doc)
                    
                    # Save processed document to bytes
                    output = io.BytesIO()
                    processed_doc.save(output)
                    output.seek(0)
                    
                    # Display results
                    st.success(f"✅ Successfully processed {tables_fixed} tables!")
                    
                    # Download button
                    original_name = os.path.splitext(uploaded_file.name)[0]
                    new_filename = f"{original_name}_fixed.docx"
                    
                    st.download_button(
                        label="📥 Download Fixed Document",
                        data=output,
                        file_name=new_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                    
                    # Show preview info
                    with st.expander("📊 Processing Summary"):
                        st.write(f"- Tables processed: **{tables_fixed}**")
                        st.write(f"- Processing mode: **{processing_mode}**")
                        st.write(f"- Original file: **{uploaded_file.name}**")
                        
            except Exception as e:
                st.error(f"❌ Error processing document: {str(e)}")
                st.info("Please make sure you've uploaded a valid Word document (.docx format)")
    
    # Instructions section
    with st.expander("ℹ️ How to use this tool"):
        st.markdown("""
        1. **Upload** your Word document using the file uploader above
        2. **Choose** between Standard or Enhanced processing mode:
           - **Standard**: Uses built-in Table Grid style (faster)
           - **Enhanced**: Applies custom borders (more control over appearance)
        3. **Click** the "Process Document" button
        4. **Download** your fixed document
        
        **What this tool fixes:**
        - Removes black/colored table backgrounds
        - Adds proper dark grid lines to tables
        - Normalizes text formatting (sets text to black, removes bold/italic from table cells)
        - Ensures consistent table appearance
        """)
    
    # Features section
    with st.expander("🔧 Features"):
        st.markdown("""
        - **Background Removal**: Converts black table backgrounds to white
        - **Grid Lines**: Adds clean, dark grid lines to tables
        - **Text Normalization**: Ensures table text is properly formatted in black
        - **Batch Processing**: Handles all tables in the document automatically
        - **Format Preservation**: Maintains document structure while fixing tables
        """)

if __name__ == "__main__":
    main()
